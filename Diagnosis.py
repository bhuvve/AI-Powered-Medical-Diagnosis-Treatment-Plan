import os
import streamlit as st
from crewai import Agent, Task, Crew, LLM
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import base64

# Load environment variables
load_dotenv()

# Set API keys from environment variables
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")

# Initialize LLM (Replace with actual instance)
llm = LLM(model="gpt-3.5-turbo")  # Adjust model if needed

# Initialize tools
search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()

# Function to generate DOCX
def generate_docx(content):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    doc.add_paragraph(content)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Function to create a download link
def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

# Streamlit UI
st.title("AI-Powered Medical Diagnosis & Treatment")
symptoms = st.text_area("Enter Patient Symptoms:")
medical_history = st.text_area("Enter Medical History:")

# Define Agents
diagnostician = Agent(
    role="Medical Diagnostician",
    goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
    backstory="Expert in diagnosing medical conditions based on patient-reported symptoms and history.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

treatment_advisor = Agent(
    role="Treatment Advisor",
    goal="Recommend appropriate treatment plans based on the diagnosis.",
    backstory="Expert in developing patient-specific treatment strategies.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

medical_testing_advisor = Agent(
    role="Medical Testing Advisor",
    goal="Suggest necessary medical tests based on diagnosis and symptoms.",
    backstory="Expert in recommending diagnostic tests such as MRI, blood tests, and CBP.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

# Define Tasks
diagnose_task = Task(
    description=f"""
        1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).
        2. Provide a preliminary diagnosis with possible conditions.
        3. Limit the diagnosis to the most likely conditions.
    """,
    expected_output="A preliminary diagnosis with possible conditions.",
    agent=diagnostician
)

treatment_task = Task(
    description=f"""
        1. Based on the diagnosis, recommend step-by-step treatment plans.
        2. Consider the patient's medical history ({medical_history}) and symptoms ({symptoms}).
        3. Provide detailed treatment recommendations including medications and lifestyle changes.
    """,
    expected_output="A comprehensive treatment plan.",
    agent=treatment_advisor
)

testing_task = Task(
    description="""
        1. Recommend necessary medical tests based on diagnosis and symptoms.
        2. Suggest tests such as MRI, blood tests, and CBP where relevant.
        3. Explain why each test is recommended and how it aids in diagnosis.
    """,
    expected_output="A list of recommended tests with justifications.",
    agent=medical_testing_advisor
)

# Create Crew
crew = Crew(
    agents=[diagnostician, treatment_advisor, medical_testing_advisor],
    tasks=[diagnose_task, treatment_task, testing_task],
    verbose=True
)

# Execution
if st.button("Get Diagnosis and Treatment Plan"):
    if symptoms and medical_history:
        with st.spinner('Generating recommendations...'):
            try:
                result = crew.kickoff()
                
                # Extract and format task outputs
                diagnosis = result.tasks_output[0].raw if result.tasks_output else "No diagnosis available."
                treatment = result.tasks_output[1].raw if result.tasks_output else "No treatment available."
                tests = result.tasks_output[2].raw if result.tasks_output else "No test recommendations available."

                formatted_result = (
                    f"## Diagnosis:\n{diagnosis}\n\n"
                    f"## Treatment Plan:\n{treatment}\n\n"
                    f"## Recommended Tests:\n{tests}"
                )

                # Display results in Streamlit
                st.write(formatted_result)

                # Generate and provide DOCX download
                docx_file = generate_docx(formatted_result)
                download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
                st.markdown(download_link, unsafe_allow_html=True)

            except Exception as e:
                st.error(f"An error occurred: {e}")
    else:
        st.warning("Please enter both symptoms and medical history.")
